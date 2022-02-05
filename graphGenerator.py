# import modules
from datetime import datetime

import docx
import matplotlib.pyplot as plt
import yfinance as yf
from docx import Document
from docx.shared import Inches
from docx2pdf import convert


def oneyear(company):
    # initialize parameters
    start_date = datetime(2021, 1, 1)
    end_date = datetime(2022, 1, 1)

    # get the data
    data = yf.download(company, start=start_date,
                       end=end_date)

    # display
    plt.figure(figsize=(20, 10))
    plt.title('Opening Prices from {} to {}'.format(start_date,
                                                    end_date))
    plt.plot(data['Open'])
    plt.savefig('1y.png')


def halfyear(company):
    start_date = datetime(2021, 6, 1)
    end_date = datetime(2022, 1, 1)

    # get the data
    data = yf.download(company, start=start_date,
                       end=end_date)

    # display
    plt.figure(figsize=(20, 10))
    plt.title('Opening Prices from {} to {}'.format(start_date,
                                                    end_date))
    plt.plot(data['Open'])
    # plt.show()
    plt.savefig('6m.png')


def threemonths(company):
    start_date = datetime(2021, 9, 1)
    end_date = datetime(2022, 1, 1)

    # get the data
    data = yf.download(company, start=start_date,
                       end=end_date)

    # display
    plt.figure(figsize=(20, 10))
    plt.title('Opening Prices from {} to {}'.format(start_date,
                                                    end_date))
    plt.plot(data['Open'])
    # plt.show()
    plt.savefig('3m.png')


def oneweek(company):
    start_date = datetime(2022, 1, 20)
    end_date = datetime(2022, 1, 29)

    # get the data
    data = yf.download(company, start=start_date,
                       end=end_date)

    # display
    plt.figure(figsize=(20, 10))
    plt.title('Opening Prices from {} to {}'.format(start_date,
                                                    end_date))
    plt.plot(data['Open'])
    # plt.show()
    plt.savefig('1d.png')


def onemonth(company):
    start_date = datetime(2022, 1, 1)
    end_date = datetime(2022, 2, 2)

    # get the data
    data = yf.download(company, start=start_date,
                       end=end_date)

    # display
    plt.figure(figsize=(20, 10))
    plt.title('Opening Prices from {} to {}'.format(start_date,
                                                    end_date))
    plt.plot(data['Open'])
    # plt.show()
    plt.savefig('1m.png')


def convertion():
    document = Document()
    document.add_heading('1  Day', 1)
    document.add_picture('1d.png', width=Inches(6))
    document.add_heading('1 Month', 1)
    document.add_picture('1m.png', width=Inches(6))
    document.add_heading('3 Months ', 1)
    document.add_picture('3m.png', width=Inches(6))
    document.add_heading('6 Months', 1)
    document.add_picture('6m.png', width=Inches(6))
    # doc = docx.Document()
    # p1 = doc.add_paragraph()
    # p1=doc.add_heading('1 year')
    # p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading('1  Year', 1)
    document.add_picture('1y.png', width=Inches(6))

    # document.add_picture('stock12M.png', width=Inches(6))

    # document.add_page_break()

    # document.add_picture('sample.jpg')

    # write to docx
    document.save('stocks.docx')

    convert("stocks.docx")


# oneyear()
# halfyear()
# threemonths()
# oneweek()
# onemonth()
# convertion()
